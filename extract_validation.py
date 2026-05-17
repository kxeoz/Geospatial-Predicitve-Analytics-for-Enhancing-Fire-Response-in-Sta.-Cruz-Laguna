from docx import Document

doc = Document('VALIDATION.docx')
print("=== DOCUMENT CONTENT ===")
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)
        
print("\n=== TABLES ===")
for i, table in enumerate(doc.tables):
    print(f"\nTable {i+1}:")
    for j, row in enumerate(table.rows):
        cells = [cell.text for cell in row.cells]
        print(f"Row {j}: {cells}")
